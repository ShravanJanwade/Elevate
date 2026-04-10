import os
import tempfile
from functools import wraps

import requests
from dotenv import load_dotenv
from flask import Flask, request, jsonify, g
from flask_cors import CORS

from resume_parser import parse_resume, parse_resume_from_pdf
from analyzer import full_analysis, section_scores, semantic_score
from suggestion_generator import generate_suggestions, rewrite_bullet, _rule_based_rewrite

load_dotenv()

app = Flask(__name__)
CORS(app)

SUPABASE_URL = os.environ.get("SUPABASE_URL", "")
SUPABASE_SERVICE_ROLE_KEY = os.environ.get("SUPABASE_SERVICE_ROLE_KEY", "")
SUPABASE_ANON_KEY = os.environ.get("SUPABASE_ANON_KEY", "")


# ---------------------------------------------------------------------------
# Auth middleware
# ---------------------------------------------------------------------------

def require_auth(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        auth_header = request.headers.get("Authorization", "")
        if not auth_header.startswith("Bearer "):
            return jsonify({"error": "Unauthorized"}), 401
        token = auth_header[7:]
        resp = requests.get(
            f"{SUPABASE_URL}/auth/v1/user",
            headers={"Authorization": f"Bearer {token}", "apikey": SUPABASE_ANON_KEY},
            timeout=10,
        )
        if resp.status_code != 200:
            return jsonify({"error": "Unauthorized"}), 401
        user_data = resp.json()
        g.user_id = user_data.get("id")
        g.token = token
        return f(*args, **kwargs)
    return decorated


def optional_auth(f):
    """Auth that doesn't fail — sets g.user_id to None if unauthenticated."""
    @wraps(f)
    def decorated(*args, **kwargs):
        auth_header = request.headers.get("Authorization", "")
        g.user_id = None
        g.token = None
        if auth_header.startswith("Bearer "):
            token = auth_header[7:]
            try:
                resp = requests.get(
                    f"{SUPABASE_URL}/auth/v1/user",
                    headers={"Authorization": f"Bearer {token}", "apikey": SUPABASE_ANON_KEY},
                    timeout=10,
                )
                if resp.status_code == 200:
                    user_data = resp.json()
                    g.user_id = user_data.get("id")
                    g.token = token
            except Exception:
                pass
        return f(*args, **kwargs)
    return decorated


# ---------------------------------------------------------------------------
# Supabase REST helpers
# ---------------------------------------------------------------------------

def _sb_headers(token=None):
    # Use user JWT if provided, otherwise fall back to service role key
    auth_token = token or SUPABASE_SERVICE_ROLE_KEY
    api_key = SUPABASE_ANON_KEY or SUPABASE_SERVICE_ROLE_KEY
    return {
        "apikey": api_key,
        "Authorization": f"Bearer {auth_token}",
        "Content-Type": "application/json",
        "Prefer": "return=representation",
    }


def _insert(table, payload, token=None):
    r = requests.post(
        f"{SUPABASE_URL}/rest/v1/{table}",
        json=payload,
        headers=_sb_headers(token),
        timeout=15,
    )
    r.raise_for_status()
    return r.json()


def _update(table, row_id, payload, token=None):
    r = requests.patch(
        f"{SUPABASE_URL}/rest/v1/{table}?id=eq.{row_id}",
        json=payload,
        headers=_sb_headers(token),
        timeout=15,
    )
    r.raise_for_status()
    return r.json()


def _select(table, params, token=None):
    r = requests.get(
        f"{SUPABASE_URL}/rest/v1/{table}",
        params=params,
        headers=_sb_headers(token),
        timeout=15,
    )
    r.raise_for_status()
    return r.json()


# ---------------------------------------------------------------------------
# Helper: parse resume from request
# ---------------------------------------------------------------------------

def _parse_resume_from_request(data, files):
    """Extract and parse resume from either file upload or text."""
    resume_file = files.get("resume_file") if files else None
    resume_text = data.get("resume_text", "") or ""

    if resume_file:
        suffix = os.path.splitext(resume_file.filename)[1]
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            resume_file.save(tmp.name)
            tmp_path = tmp.name
        try:
            if suffix.lower() == ".pdf":
                resume_sections = parse_resume_from_pdf(tmp_path)
            elif suffix.lower() == ".docx":
                from docx import Document as DocxDocument
                doc = DocxDocument(tmp_path)
                raw = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                resume_sections = parse_resume(raw)
            else:  # .txt and anything else
                with open(tmp_path, "r", encoding="utf-8", errors="ignore") as f:
                    raw = f.read()
                resume_sections = parse_resume(raw)
        finally:
            os.unlink(tmp_path)
    elif resume_text:
        resume_sections = parse_resume(resume_text)
    else:
        return None

    return resume_sections


# ---------------------------------------------------------------------------
# Health
# ---------------------------------------------------------------------------

@app.route("/api/health", methods=["GET"])
def health():
    return jsonify({"status": "ok"})


# ---------------------------------------------------------------------------
# POST /api/preview-pdf  — convert .docx / .txt to PDF for preview
# ---------------------------------------------------------------------------

@app.route("/api/preview-pdf", methods=["POST"])
def preview_pdf():
    """Convert a .docx or .txt file to PDF and return it as application/pdf."""
    from flask import send_file
    import io

    file = request.files.get("file")
    if not file:
        return jsonify({"error": "No file provided"}), 400

    suffix = os.path.splitext(file.filename)[1].lower()
    if suffix not in (".docx", ".txt"):
        return jsonify({"error": "Only .docx and .txt files are supported"}), 400

    # Extract text
    try:
        if suffix == ".docx":
            from docx import Document as DocxDocument
            # On Windows, close the temp file before writing to it to avoid locking
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            tmp_path = tmp.name
            tmp.close()
            file.save(tmp_path)
            try:
                doc = DocxDocument(tmp_path)
                paragraphs = []
                for p in doc.paragraphs:
                    txt = p.text.strip()
                    if txt:
                        paragraphs.append(txt)
                raw_text = "\n".join(paragraphs)
            finally:
                os.unlink(tmp_path)
        else:  # .txt
            raw_text = file.read().decode("utf-8", errors="ignore")
    except Exception as exc:
        return jsonify({"error": f"Could not read file: {exc}"}), 500

    if not raw_text.strip():
        return jsonify({"error": "No readable text found in file"}), 400

    # Build PDF with fpdf2
    try:
        from fpdf import FPDF

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Helvetica", size=10)
        pdf.set_left_margin(15)
        pdf.set_right_margin(15)

        for line in raw_text.split("\n"):
            # Strip to ASCII-safe chars — fpdf core fonts are latin-1
            safe_line = line.encode("latin-1", errors="replace").decode("latin-1")
            pdf.multi_cell(0, 5.5, safe_line)

        pdf_bytes = bytes(pdf.output())
        buf = io.BytesIO(pdf_bytes)
        buf.seek(0)
        return send_file(buf, mimetype="application/pdf", as_attachment=False,
                         download_name="preview.pdf")
    except Exception as exc:
        return jsonify({"error": f"PDF generation failed: {exc}"}), 500


# ---------------------------------------------------------------------------
# POST /api/analyze  (authenticated)
# ---------------------------------------------------------------------------

@app.route("/api/analyze", methods=["POST"])
@require_auth
def analyze():
    data = request.get_json(silent=True) or {}
    jd_text = data.get("job_description", "") or request.form.get("job_description", "")
    data["resume_text"] = data.get("resume_text", "") or request.form.get("resume_text", "")

    if not jd_text:
        return jsonify({"error": "Job description is required."}), 400

    resume_sections = _parse_resume_from_request(data, request.files)
    if resume_sections is None:
        return jsonify({"error": "Please provide a resume (file or text)."}), 400

    try:
        analysis = full_analysis(resume_sections, jd_text)
        suggestions = generate_suggestions(
            resume_sections, jd_text, analysis["semantic_analysis"], max_suggestions=5
        )
        analysis["suggestions"] = suggestions

        # Include entities if available
        entities = resume_sections.get("entities", {})
        if entities:
            analysis["entities"] = entities

        # Save session to Supabase (best-effort)
        session_id = None
        try:
            session_rows = _insert("analysis_sessions", {
                "user_id": g.user_id,
                "job_description": jd_text,
                "resume_text": data.get("resume_text", ""),
                "overall_score": analysis["overall_score"],
                "keyword_score": analysis["keyword_analysis"]["match_percentage"],
                "semantic_score": analysis["semantic_analysis"]["overall_score"],
                "missing_keywords": analysis["keyword_analysis"]["missing"],
                "matched_keywords": analysis["keyword_analysis"]["matched"],
            }, token=g.token)
            session_id = session_rows[0]["id"] if session_rows else None

            if session_id:
                bullets = analysis["semantic_analysis"].get("bullet_scores", [])
                for b in bullets:
                    _insert("bullet_scores", {
                        "session_id": session_id,
                        "bullet_text": b["text"],
                        "similarity_score": b["similarity"],
                    }, token=g.token)
        except Exception:
            pass  # DB save failed silently; analysis results still returned

        analysis["session_id"] = session_id
        return jsonify(analysis)

    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ---------------------------------------------------------------------------
# POST /api/analyze-guest  (no auth required — for demo / quick analysis)
# ---------------------------------------------------------------------------

@app.route("/api/analyze-guest", methods=["POST"])
@optional_auth
def analyze_guest():
    """Run analysis without requiring authentication."""
    data = request.get_json(silent=True) or {}
    jd_text = data.get("job_description", "") or request.form.get("job_description", "")
    data["resume_text"] = data.get("resume_text", "") or request.form.get("resume_text", "")

    if not jd_text:
        return jsonify({"error": "Job description is required."}), 400

    resume_sections = _parse_resume_from_request(data, request.files)
    if resume_sections is None:
        return jsonify({"error": "Please provide a resume (file or text)."}), 400

    try:
        analysis = full_analysis(resume_sections, jd_text)
        suggestions = generate_suggestions(
            resume_sections, jd_text, analysis["semantic_analysis"], max_suggestions=5
        )
        analysis["suggestions"] = suggestions

        entities = resume_sections.get("entities", {})
        if entities:
            analysis["entities"] = entities

        # Try to save if authenticated
        session_id = None
        if g.user_id:
            try:
                session_rows = _insert("analysis_sessions", {
                    "user_id": g.user_id,
                    "job_description": jd_text,
                    "resume_text": data.get("resume_text", ""),
                    "overall_score": analysis["overall_score"],
                    "keyword_score": analysis["keyword_analysis"]["match_percentage"],
                    "semantic_score": analysis["semantic_analysis"]["overall_score"],
                    "missing_keywords": analysis["keyword_analysis"]["missing"],
                    "matched_keywords": analysis["keyword_analysis"]["matched"],
                }, token=g.token)
                session_id = session_rows[0]["id"] if session_rows else None
            except Exception:
                pass

        analysis["session_id"] = session_id
        return jsonify(analysis)

    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ---------------------------------------------------------------------------
# POST /api/analyze/sections
# ---------------------------------------------------------------------------

@app.route("/api/analyze/sections", methods=["POST"])
@require_auth
def analyze_sections():
    data = request.get_json(silent=True) or {}
    jd_text = data.get("job_description", "") or request.form.get("job_description", "")
    data["resume_text"] = data.get("resume_text", "") or request.form.get("resume_text", "")
    session_id = data.get("session_id")

    if not jd_text:
        return jsonify({"error": "Job description is required."}), 400

    resume_sections = _parse_resume_from_request(data, request.files)
    if resume_sections is None:
        return jsonify({"error": "Please provide a resume (file or text)."}), 400

    try:
        scores = section_scores(resume_sections, jd_text)

        # Save to section_scores table if session_id provided
        if session_id:
            for s in scores:
                sec_name = s["section"]
                sec_text = resume_sections.get(sec_name, s.get("preview", ""))
                _insert("section_scores", {
                    "session_id": session_id,
                    "section_name": sec_name,
                    "section_text": sec_text,
                    "similarity_score": s["score"],
                }, token=g.token)

        return jsonify({
            "sections": [
                {
                    "name": s["section"],
                    "score": s["score"],
                    "strength": s.get("strength", "moderate"),
                    "weight": s.get("weight", 0.5),
                }
                for s in scores
            ]
        })

    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ---------------------------------------------------------------------------
# POST /api/rescore/bullet
# ---------------------------------------------------------------------------

@app.route("/api/rescore/bullet", methods=["POST"])
@require_auth
def rescore_bullet():
    data = request.get_json(silent=True) or {}
    bullet_text = data.get("bullet", "")
    jd_text = data.get("job_description", "")
    bullet_id = data.get("bullet_id", "")

    if not bullet_text or not jd_text:
        return jsonify({"error": "bullet and job_description are required."}), 400

    try:
        from models.semantic_engine import (
            _get_bi_encoder, _get_cross_encoder,
            calibrate_semantic, score_to_strength, _normalize_cross_score,
        )
        import numpy as np
        from sklearn.metrics.pairwise import cosine_similarity as cos_sim

        bi_enc = _get_bi_encoder()
        cross_enc = _get_cross_encoder()

        jd_emb = bi_enc.encode([jd_text], show_progress_bar=False)
        b_emb = bi_enc.encode([bullet_text], show_progress_bar=False)
        bi_score = float(cos_sim(b_emb, jd_emb).flatten()[0])

        cross_raw = float(cross_enc.predict([(bullet_text, jd_text)])[0])
        cross_score = _normalize_cross_score(cross_raw)

        combined = 0.65 * bi_score + 0.35 * cross_score
        calibrated = round(calibrate_semantic(combined), 1)
        strength = score_to_strength(calibrated)

        if bullet_id:
            _update("bullet_scores", bullet_id, {
                "bullet_text": bullet_text,
                "similarity_score": calibrated,
                "rewritten_text": bullet_text,
            }, token=g.token)

        return jsonify({"new_score": calibrated, "strength": strength})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ---------------------------------------------------------------------------
# POST /api/rewrite
# ---------------------------------------------------------------------------

@app.route("/api/rewrite", methods=["POST"])
@require_auth
def rewrite():
    data = request.get_json(silent=True) or {}
    bullet_text = data.get("bullet", "")
    jd_text = data.get("job_description", "")
    use_llm = data.get("use_llm", False)
    bullet_id = data.get("bullet_id", "")

    if not bullet_text or not jd_text:
        return jsonify({"error": "bullet and job_description are required."}), 400

    try:
        if use_llm and os.environ.get("ELEVATE_USE_LLM", "false").lower() == "true":
            result = rewrite_bullet(bullet_text, jd_text)
            rewritten = result["improved"]
            method = result.get("method", "flan_t5")
            issues = result.get("issues", [])
        else:
            rewritten = _rule_based_rewrite(bullet_text, jd_text)
            method = "rule_based"
            issues = []

        if bullet_id:
            _update("bullet_scores", bullet_id, {
                "rewritten_text": rewritten,
                "rewrite_method": method,
            }, token=g.token)

        return jsonify({"rewritten": rewritten, "method": method, "issues": issues})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ---------------------------------------------------------------------------
# GET /api/history
# ---------------------------------------------------------------------------

@app.route("/api/history", methods=["GET"])
@require_auth
def history():
    try:
        rows = _select("analysis_sessions", {
            "user_id": f"eq.{g.user_id}",
            "order": "created_at.desc",
            "limit": 20,
            "select": "id,created_at,overall_score,session_label,job_description",
        }, token=g.token)
        sessions = []
        for r in rows:
            sessions.append({
                "id": r["id"],
                "created_at": r["created_at"],
                "overall_score": r["overall_score"],
                "session_label": r.get("session_label"),
                "job_description_preview": (r.get("job_description") or "")[:100],
            })
        return jsonify({"sessions": sessions})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ---------------------------------------------------------------------------
# GET /api/history/<session_id>
# ---------------------------------------------------------------------------

@app.route("/api/history/<session_id>", methods=["GET"])
@require_auth
def history_detail(session_id):
    try:
        sessions = _select("analysis_sessions", {
            "id": f"eq.{session_id}",
            "user_id": f"eq.{g.user_id}",
        }, token=g.token)
        if not sessions:
            return jsonify({"error": "Session not found"}), 404

        session = sessions[0]

        bullets = _select("bullet_scores", {"session_id": f"eq.{session_id}"}, token=g.token)
        sections = _select("section_scores", {"session_id": f"eq.{session_id}"}, token=g.token)

        return jsonify({
            "session": session,
            "bullets": bullets,
            "section_scores": sections,
        })
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


if __name__ == "__main__":
    app.run(debug=True, port=5000)
# updated routes
# added CORS
# update imports
# api keys
# final endpoints
# route patch
